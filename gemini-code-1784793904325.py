import streamlit as st
import os
import tempfile
import subprocess
import whisper
from docx import Document
from googletrans import Translator

# Configurazione della pagina Streamlit
st.set_page_config(page_title="Video Subtitler & Transcriber", page_icon="🎬", layout="centered")

st.title("🎬 Video Subtitler & Transcriber")
st.write("Carica o trascina un video per aggiungere i sottotitoli in stile film e generare il documento Word in Italiano e Inglese.")

# 1. Componente Drag & Drop per il caricamento del video
uploaded_file = st.file_uploader("Trascina o seleziona il tuo video qui", type=["mp4", "mov", "avi", "mkv"])

def format_timestamp(seconds):
    hrs = int(seconds // 3600)
    mins = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    msecs = int((seconds - int(seconds)) * 1000)
    return f"{hrs:02d}:{mins:02d}:{secs:02d},{msecs:03d}"

if uploaded_file is not None:
    st.video(uploaded_file)
    
    if st.button("🚀 Avvia Elaborazione Video e Trascrizione"):
        with st.spinner("Elaborazione in corso... Questo processo può richiedere qualche minuto."):
            try:
                # Salva il file caricato in una cartella temporanea
                with tempfile.TemporaryDirectory() as temp_dir:
                    input_video_path = os.path.join(temp_dir, uploaded_file.name)
                    with open(input_video_path, "wb") as f:
                        f.write(uploaded_file.read())

                    base_name = os.path.splitext(uploaded_file.name)[0]
                    output_video_path = os.path.join(temp_dir, f"{base_name}_subtitled.mp4")
                    output_docx_path = os.path.join(temp_dir, f"{base_name}_trascrizione.docx")
                    srt_path = os.path.join(temp_dir, "subtitles.srt")

                    # FASE 1: Trascrizione AI con Whisper
                    st.info("1/4 Trascrizione dell'audio con AI in corso...")
                    model = whisper.load_model("base")
                    result = model.transcribe(input_video_path, language="it")

                    # FASE 2: Creazione File SRT (Sottotitoli)
                    st.info("2/4 Generazione sottotitoli...")
                    with open(srt_path, "w", encoding="utf-8") as f:
                        for i, segment in enumerate(result["segments"], start=1):
                            start = format_timestamp(segment["start"])
                            end = format_timestamp(segment["end"])
                            text = segment["text"].strip()
                            f.write(f"{i}\n{start} --> {end}\n{text}\n\n")

                    # FASE 3: Hardcoding Sottotitoli sul Video (Stile Film)
                    st.info("3/4 Montaggio sottotitoli in stile film sul video...")
                    clean_srt_path = srt_path.replace("\\", "/").replace(":", "\\:")
                    style = "Fontname=Arial,Fontsize=18,PrimaryColour=&H00FFFFFF&,BackColour=&H80000000&,BorderStyle=4,MarginV=30"
                    
                    ffmpeg_cmd = [
                        "ffmpeg", "-y", "-i", input_video_path,
                        "-vf", f"subtitles='{clean_srt_path}':force_style='{style}'",
                        "-c:a", "copy", output_video_path
                    ]
                    subprocess.run(ffmpeg_cmd, check=True)

                    # FASE 4: Documento Word (IT + EN)
                    st.info("4/4 Traduzione in Inglese e generazione file Word...")
                    doc = Document()
                    doc.add_heading(f"Trascrizione: {base_name}", level=1)

                    # Italiano
                    doc.add_heading("Testo Originale (Italiano)", level=2)
                    full_text_it = result["text"].strip()
                    doc.add_paragraph(full_text_it)

                    # Traduzione Inglese
                    translator = Translator()
                    translated_text = translator.translate(full_text_it, src="it", dest="en").text

                    doc.add_heading("English Translation", level=2)
                    doc.add_paragraph(translated_text)

                    doc.save(output_docx_path)

                    st.success("✅ Elaborazione completata con successo!")

                    # Pulsanti di Download per l'utente finale
                    with open(output_video_path, "rb") as v_file:
                        st.download_button(
                            label="📥 Scarica Video Sottotitolato",
                            data=v_file,
                            file_name=f"{base_name}_subtitled.mp4",
                            mime="video/mp4"
                        )

                    with open(output_docx_path, "rb") as d_file:
                        st.download_button(
                            label="📥 Scarica Documento Word (IT + EN)",
                            data=d_file,
                            file_name=f"{base_name}_trascrizione.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            except Exception as e:
                st.error(f"Si è verificato un errore durante l'elaborazione: {e}")